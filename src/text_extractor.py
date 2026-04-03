import logging
from typing import Dict, Any
from src.ocr_provider import OCRFactory

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class TextExtractor:
    """Text extractor using configured OCR provider"""
    
    def __init__(self):
        self.ocr_provider = OCRFactory.get_provider()
        logger.info(f"Using OCR: {self.ocr_provider.get_provider_name()}")
    
    def extract_from_pdf(self, file_content: bytes) -> str:
        """Extract text from PDF"""
        try:
            import PyPDF2
            import io
            
            text = ""
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
            
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n\n"
            
            if not text.strip():
                # Try OCR for scanned PDF
                text = self._ocr_pdf(file_content)
            
            return text.strip() if text.strip() else "No readable text found in PDF."
            
        except Exception as e:
            logger.error(f"PDF extraction error: {str(e)}")
            return f"Error extracting PDF: {str(e)}"
    
    def _ocr_pdf(self, file_content: bytes) -> str:
        """OCR for scanned PDFs"""
        try:
            from pdf2image import convert_from_bytes
            
            images = convert_from_bytes(file_content, dpi=200)
            text = ""
            
            for i, image in enumerate(images):
                # Convert image to bytes
                import io
                img_bytes = io.BytesIO()
                image.save(img_bytes, format='PNG')
                img_bytes = img_bytes.getvalue()
                
                # Use OCR provider
                page_text = self.ocr_provider.extract_text(img_bytes)
                text += f"Page {i+1}:\n{page_text}\n\n"
            
            return text.strip()
            
        except Exception as e:
            logger.error(f"PDF OCR error: {str(e)}")
            return ""
    
    def extract_from_docx(self, file_content: bytes) -> str:
        """Extract text from DOCX"""
        try:
            from docx import Document
            import io
            
            doc = Document(io.BytesIO(file_content))
            text_parts = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text.strip())
            
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                    if row_text:
                        text_parts.append(row_text)
            
            return "\n\n".join(text_parts) if text_parts else "No text found in DOCX file."
            
        except Exception as e:
            logger.error(f"DOCX extraction error: {str(e)}")
            return f"Error extracting DOCX: {str(e)}"
    
    def extract_from_image(self, file_content: bytes) -> str:
        """Extract text from image using configured OCR provider"""
        try:
            text = self.ocr_provider.extract_text(file_content)
            return text if text else "No text detected in image."
            
        except Exception as e:
            logger.error(f"Image extraction error: {str(e)}")
            return f"Error processing image: {str(e)}"
    
    def extract_text(self, file_content: bytes, file_type: str) -> str:
        """Main extraction method"""
        logger.info(f"Extracting from {file_type}")
        
        if file_type.lower() == "pdf":
            return self.extract_from_pdf(file_content)
        elif file_type.lower() == "docx":
            return self.extract_from_docx(file_content)
        elif file_type.lower() == "image":
            return self.extract_from_image(file_content)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")